"""Converts Markdown with LaTeX-like formulas to Word documents."""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class MarkdownToWordConverter:
    """Converts Markdown with LaTeX-like formulas to Word with Open Sans font."""

    FONT_NAME = "Open Sans"
    FONT_SIZE_PT = 18

    @staticmethod
    def latex_like_to_runs(content: str):
        runs = []
        i, n = 0, len(content)
        buf: list[str] = []

        def flush_buf():
            if buf:
                runs.append(("".join(buf), False, False))
                buf.clear()

        while i < n:
            char = content[i]
            if char in "_^":
                is_sub = char == "_"
                i += 1
                if i < n and content[i] == "{":
                    closing = content.find("}", i + 1)
                    if closing == -1:
                        buf.append(char)
                        continue
                    inner = content[i + 1 : closing]
                    flush_buf()
                    runs.append((inner, is_sub, not is_sub))
                    i = closing + 1
                    continue
                if i < n:
                    ch = content[i]
                    flush_buf()
                    runs.append((ch, is_sub, not is_sub))
                    i += 1
                    continue
                buf.append(char)
            else:
                buf.append(char)
                i += 1
        flush_buf()
        return runs

    @staticmethod
    def convert_latex_to_simple(latex_content: str) -> str:
        converted = latex_content
        # Keep content of common text commands
        converted = re.sub(r"\\mathrm\{([^}]+)\}", r"\1", converted)
        converted = re.sub(r"\\text\{([^}]+)\}", r"\1", converted)
        # Arrows and common symbols
        converted = re.sub(r"\\uparrow", "↑", converted)
        converted = re.sub(r"\\downarrow", "↓", converted)
        converted = re.sub(r"\\to|\\longrightarrow|-{1,2}>", "→", converted)
        converted = re.sub(r"\\leftrightarrow|<->|<=>", "↔", converted)
        converted = re.sub(r"\\xrightarrow\[[^]]*\]\{[^}]*\}", "→", converted)
        # Remove begin/end and alignment helpers
        converted = re.sub(r"\\begin\{[^}]+\}|\\end\{[^}]+\}", "", converted)
        # LaTeX line breaks → newline
        converted = re.sub(r"\\\\", "\n", converted)
        # Remove alignment ampersands
        converted = re.sub(r"\s*&\s*", " ", converted)
        # OCR artifacts like { i }{=} or { I }{=} → '='
        converted = re.sub(r"\{\s*[^}]*\s*\}\s*\{=\}", " = ", converted)
        # Keep content of remaining simple commands
        converted = re.sub(r"\\[a-zA-Z]+\*?\{([^}]*)\}", r"\1", converted)
        # Drop other commands
        converted = re.sub(r"\\[a-zA-Z]+\*?", "", converted)
        # Remove stray braces
        converted = converted.replace("{", "").replace("}", "")
        # Normalize spaces per line, keep newlines
        converted = re.sub(r"[\t ]+", " ", converted)
        converted = re.sub(r" *\n *", "\n", converted)
        converted = converted.strip()
        return converted

    @staticmethod
    def parse_content(text: str):
        """Split markdown into text, inline ($...$) and block ($$...$$) formulas.
        Block formulas may contain multiple lines (\\ in LaTeX). We split such
        blocks into individual lines so that each renders on its own centered line.
        """
        segments: list[tuple[str, str]] = []
        parts = re.split(r"(\$\$[\s\S]*?\$\$)", text)

        for part in parts:
            if not part:
                continue
            if part.startswith("$$") and part.endswith("$$"):
                latex_content = part[2:-2].strip()
                converted = MarkdownToWordConverter.convert_latex_to_simple(latex_content)
                for line in converted.split("\n"):
                    if line.strip():
                        segments.append(("block_formula", line.strip()))
            else:
                inline_parts = re.split(r"(\$[^$]+\$)", part)
                for inline in inline_parts:
                    if not inline:
                        continue
                    if inline.startswith("$") and inline.endswith("$"):
                        latex_content = inline[1:-1]
                        converted = MarkdownToWordConverter.convert_latex_to_simple(latex_content)
                        segments.append(("inline_formula", converted))
                    elif inline.strip():
                        segments.append(("text", inline))
        return segments

    @staticmethod
    def normalize_chem_spacing(text: str) -> str:
        """Fix typical OCR spacing in chemical formulas (H 2 O → H2O, Si Cl4 → SiCl4)."""
        s = text
        # Letter + spaces + digit → join
        s = re.sub(r"(?<=[A-Za-z])\s+(?=\d)", "", s)
        # Digit + spaces + letter/paren → join
        s = re.sub(r"(?<=\d)\s+(?=[A-Za-z(\[\{])", "", s)
        # Lowercase + spaces + Uppercase → join (Si Cl → SiCl)
        s = re.sub(r"(?<=[a-z])\s+(?=[A-Z])", "", s)
        # Right bracket + spaces + number (] 6 → ]6)
        s = re.sub(r"(?<=[)\]\}])\s+(?=\d)", "", s)
        # Collapse multiple spaces
        s = re.sub(r"\s+", " ", s).strip()
        return s

    @staticmethod
    def _split_for_subscripts(text: str):
        """Split so that digits following a letter or closing bracket become subscript runs."""
        runs = []
        i = 0
        n = len(text)
        while i < n:
            if text[i].isdigit() and i > 0 and (text[i - 1].isalpha() or text[i - 1] in ")]}"):
                j = i
                while j < n and text[j].isdigit():
                    j += 1
                runs.append((text[i:j], True, False))
                i = j
                continue
            start = i
            i += 1
            while i < n and not (text[i].isdigit() and i > 0 and (text[i - 1].isalpha() or text[i - 1] in ")]}")):
                i += 1
            runs.append((text[start:i], False, False))
        return runs

    def _make_runs(self, content: str):
        """Combine LaTeX ^/_ parsing with chemical subscript heuristics."""
        normalized = self.normalize_chem_spacing(content)
        base_runs = self.latex_like_to_runs(normalized)
        final_runs = []
        for text, is_sub, is_sup in base_runs:
            if is_sub or is_sup:
                final_runs.append((text, is_sub, is_sup))
            else:
                final_runs.extend(self._split_for_subscripts(text))
        return final_runs

    def _apply_run_style(self, run):
        run.font.name = self.FONT_NAME
        run.font.size = Pt(self.FONT_SIZE_PT)

    def create_word_document(self, markdown_text: str) -> bytes:
        segments = self.parse_content(markdown_text)
        document = Document()

        try:
            normal_style = document.styles["Normal"]
        except KeyError:
            normal_style = None

        if normal_style is not None:
            normal_style.font.name = self.FONT_NAME
            normal_style.font.size = Pt(self.FONT_SIZE_PT)

        current_paragraph = None

        for segment_type, content in segments:
            if segment_type == "block_formula":
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for text, is_sub, is_sup in self._make_runs(content):
                    run = paragraph.add_run(text)
                    self._apply_run_style(run)
                    if is_sub:
                        run.font.subscript = True
                    if is_sup:
                        run.font.superscript = True
                current_paragraph = None
            elif segment_type == "inline_formula":
                if current_paragraph is None:
                    current_paragraph = document.add_paragraph()
                for text, is_sub, is_sup in self._make_runs(content):
                    run = current_paragraph.add_run(text)
                    self._apply_run_style(run)
                    if is_sub:
                        run.font.subscript = True
                    if is_sup:
                        run.font.superscript = True
            else:
                lines = content.split("\n")
                for index, line in enumerate(lines):
                    line = line.strip()
                    if line:
                        if current_paragraph is None:
                            current_paragraph = document.add_paragraph()
                        for text, is_sub, is_sup in self._make_runs(line):
                            run = current_paragraph.add_run(text)
                            self._apply_run_style(run)
                            if is_sub:
                                run.font.subscript = True
                            if is_sup:
                                run.font.superscript = True
                    if index < len(lines) - 1 and line:
                        current_paragraph = None

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.read()
