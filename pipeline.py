"""Baseline-aligned chemistry pipeline.

This module rewrites the pipeline to mirror the behavior of baseline-pipeline.py:
- Use the exact ViT and text-refiner prompts from the baseline
- Perform ViT → text refinement (with retry/backoff) → minimal cleanup
- Build DOCX using the baseline's custom parser for '^' and '_' indices
- No additional heuristic postprocessing beyond baseline's steps
"""

from __future__ import annotations

import base64
import re
import time
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt
from exceptions import ChemistryOCRError
from mistralai import ImageURLChunk, Mistral, TextChunk
from ocr import ChemistryOCR

try:  # Prefer SDKError for 429 handling if available
    from mistralai.models import SDKError  # type: ignore
except Exception:  # noqa: BLE001
    SDKError = Exception  # type: ignore


# Prompts copied from baseline-pipeline.py (keep content in sync)
SYSTEM_PROMPT_CHEMISTRY_VIT = (
    "You are a careful OCR assistant for chemistry textbooks and lab notes. "
    "Return clean Markdown only. STRICT FORMAT: \n"
    "- One reaction per line. Use: $REACTANTS → PRODUCTS$ optionally followed by a space and 'conditions: ...'.\n"
    "- Examples: \n  $I_{2} + 5Cl_{2} + 6H_{2}O = 2HIO_{3} + 10HCl $\n  $ 6NaOH + 3I_{2} → NaIO_{3} + 5NaI + 3H_{2}O$ conditions: 0°C\n  $1/2 I_{2} + (CO)_{4}FeI_{2} → FeI_{3} + 4CO$ conditions: hν, hexane\n"
    "- Put ANY reaction conditions (catalyst, temperature, solvent, light hν, Δ, pressure, etc.) AFTER the reaction as 'conditions: ...'. Never place them on the arrow or inside the equation.\n"
    "- No LaTeX, which can't be rendered WITHOUT Word Equation objects. Only simple letex for for upper and lower indexes (H_{2}O, N^{2+}, C_{2}H_{5}OH).\n"
    "- If the image has multiple columns, transcribe columns SEQUENTIALLY: finish the left column top→bottom, then the right column. Do NOT mix columns.\n"
    "- Do NOT include ```markdown or ``` code fences.\n"
    "- After each reaction line output a literal \\n (backslash-n) to mark a new line."
)

SYSTEM_PROMPT_TEXT_REFINER = (
    "You are an editor for chemistry reactions. Split any glued reactions into separate lines.\n"
    "Rules:\n"
    "- Preserve all content and order; do not invent or drop tokens.\n"
    "- Exactly one reaction per line: $REACTANTS → PRODUCTS$ optionally followed by ' conditions: ...'.\n"
    "- If a complete reaction is immediately followed by another, insert a newline between them.\n"
    "- End every line with a literal \\n (backslash-n)."
)


class _BaselineDocxWriter:
    """Baseline DOCX writer that renders '^' and '_' segments as super/subscript.

    This closely follows the save_docx_from_text section of baseline-pipeline.py,
    but returns bytes instead of writing to disk.
    """

    DOC_FONT_NAME = "Open Sans"
    DOC_FONT_SIZE_PT = 18

    @staticmethod
    def _find_matching_brace(s: str, start: int) -> int:
        depth = 0
        for i in range(start, len(s)):
            if s[i] == '{':
                depth += 1
            elif s[i] == '}':
                depth -= 1
                if depth == 0:
                    return i
        return -1

    @staticmethod
    def _add_run(par, text: str, superscript: bool = False, subscript: bool = False) -> None:
        if not text:
            return
        run = par.add_run(text)
        run.font.name = _BaselineDocxWriter.DOC_FONT_NAME
        run.font.size = Pt(_BaselineDocxWriter.DOC_FONT_SIZE_PT)
        if superscript:
            run.font.superscript = True
        if subscript:
            run.font.subscript = True

    @classmethod
    def _add_parsed_line(cls, par, line: str) -> None:
        i, n = 0, len(line)
        normal: list[str] = []

        def flush() -> None:
            if normal:
                cls._add_run(par, ''.join(normal))
                normal.clear()

        while i < n:
            ch = line[i]
            if ch in ('^', '_'):
                is_super = ch == '^'
                i += 1
                content = ''

                if i < n and line[i] == '{':
                    j = cls._find_matching_brace(line, i)
                    if j == -1:
                        normal.append(ch)
                        continue
                    content = line[i+1:j]
                    i = j + 1
                else:
                    if is_super:
                        m = re.match(r'[+\-±∓]?\d+', line[i:])
                        if m:
                            content = m.group(0)
                            i += len(content)
                        else:
                            content = line[i]
                            i += 1
                    else:
                        m = re.match(r'\d+', line[i:])
                        if m:
                            content = m.group(0)
                            i += len(content)
                        else:
                            content = line[i]
                            i += 1

                flush()
                cls._add_run(par, content, superscript=is_super, subscript=not is_super)
            else:
                normal.append(ch)
                i += 1

        flush()

    @classmethod
    def create_docx_bytes(cls, text: str) -> bytes:
        doc = Document()
        try:
            normal_style = doc.styles["Normal"]
            normal_style.font.name = cls.DOC_FONT_NAME
            normal_style.font.size = Pt(cls.DOC_FONT_SIZE_PT)
        except Exception:  # noqa: BLE001
            pass

        for raw_line in text.split('\n'):
            par = doc.add_paragraph()
            if raw_line:
                cls._add_parsed_line(par, raw_line)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()


class ChemistryPipeline:
    """Full pipeline from image bytes to DOCX using baseline behavior."""

    def __init__(self, api_key: str):
        if not api_key:
            raise ChemistryOCRError("API key is not configured")
        self.api_key = api_key
        self.client = Mistral(api_key=api_key)
        self.ocr = ChemistryOCR(api_key)

    @staticmethod
    def _to_data_url(data: bytes, mime_type: str) -> str:
        return f"data:{mime_type};base64,{base64.b64encode(data).decode()}"

    def _vit_extract_text(self, data: bytes, mime_type: str, model: str) -> str:
        if not data:
            raise ChemistryOCRError("Empty image payload")

        image_data_url = self._to_data_url(data, mime_type)
        try:
            vit_response = self.client.chat.complete(
                model=model,
                messages=[
                    {"role": "system", "content": [TextChunk(text=SYSTEM_PROMPT_CHEMISTRY_VIT)]},
                    {
                        "role": "user",
                        "content": [
                            TextChunk(text="Extract Markdown as instructed."),
                            ImageURLChunk(image_url=image_data_url),
                        ],
                    },
                ],
            )
        except Exception as exc:  # noqa: BLE001
            raise ChemistryOCRError(f"Vision request failed: {exc}") from exc

        text: str | None = None
        try:
            if hasattr(vit_response, "choices") and vit_response.choices:
                msg = getattr(vit_response.choices[0], "message", None)
                if msg is not None:
                    content = getattr(msg, "content", None)
                    if isinstance(content, str):
                        text = content
                    elif isinstance(content, list):
                        parts: list[str] = []
                        for c in content:
                            chunk_text = getattr(c, "text", None)
                            if chunk_text:
                                parts.append(chunk_text)
                        if parts:
                            text = "\n".join(parts)
            if text is None and hasattr(vit_response, "output_text"):
                text = getattr(vit_response, "output_text")
        except Exception:  # noqa: BLE001
            text = None

        if not text:
            text = str(vit_response)
        return text

    def _refine_text_with_retry(self, raw_text: str, model_name: str = "mistral-small-latest", max_retries: int = 3, base_sleep: float = 1.0) -> str:
        def call_once(name: str):
            return self.client.chat.complete(
                model=name,
                messages=[
                    {"role": "system", "content": [TextChunk(text=SYSTEM_PROMPT_TEXT_REFINER)]},
                    {"role": "user", "content": [TextChunk(text=raw_text)]},
                ],
            )

        last_err: Exception | None = None
        for attempt in range(max_retries):
            try:
                resp = call_once(model_name)
                out: str | None = None
                if hasattr(resp, "choices") and resp.choices:
                    msg = getattr(resp.choices[0], "message", None)
                    if msg is not None:
                        content = getattr(msg, "content", None)
                        if isinstance(content, str):
                            out = content
                        elif isinstance(content, list):
                            segs = [getattr(c, "text", "") for c in content if getattr(c, "text", None)]
                            if segs:
                                out = "\n".join(segs)
                if not out and hasattr(resp, "output_text"):
                    out = getattr(resp, "output_text")
                return (out or str(resp))
            except SDKError as e:  # type: ignore[misc]
                status = getattr(e, "status_code", None)
                msg = str(e)
                if status == 429 or "Status 429" in msg or "capacity" in msg.lower():
                    sleep_s = base_sleep * (2 ** attempt)
                    time.sleep(sleep_s)
                    last_err = e
                    continue
                raise
            except Exception as e:  # noqa: BLE001
                last_err = e
                break
        if last_err:
            raise last_err
        return raw_text

    @staticmethod
    def _final_cleanup(text: str) -> str:
        # Convert literal "\\n" to real newlines and strip '$'
        cleaned = text.replace('\\n', '\n').replace('$', '')
        return cleaned

    def process_image_bytes(self, data: bytes, mime_type: str = "image/png") -> bytes:
        """OCR path → markdown → baseline DOCX renderer."""
        markdown_text = self.ocr.process_image_bytes(data, mime_type=mime_type, output_format="markdown")
        cleaned = self._final_cleanup(markdown_text)
        return _BaselineDocxWriter.create_docx_bytes(cleaned)

    def process_image_bytes_vit(self, data: bytes, mime_type: str = "image/png", model: str = "pixtral-12b") -> bytes:
        """ViT path matching baseline behavior, with small text refiner for both modes."""
        selected_model = model
        try:
            draft_text = self._vit_extract_text(data, mime_type=mime_type, model=selected_model)
        except ChemistryOCRError:
            if selected_model != "pixtral-12b":
                draft_text = self._vit_extract_text(data, mime_type=mime_type, model="pixtral-12b")
            else:
                raise

        try:
            refined_text = self._refine_text_with_retry(draft_text, model_name="mistral-small-latest")
        except Exception:
            refined_text = draft_text

        cleaned = self._final_cleanup(refined_text)
        return _BaselineDocxWriter.create_docx_bytes(cleaned)
